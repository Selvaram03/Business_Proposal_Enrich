import re, os, tempfile
from io import BytesIO
import pandas as pd
from docx import Document
from lxml import etree

# Core replace functions adapted from your code
def _replace_in_xml(doc_part, param_dict):
    try:
        if doc_part.element is None:
            return
        root = doc_part.element.getroottree()
    except AttributeError:
        return
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main', 
                  'v': 'urn:schemas-microsoft-com:vml'}
    for key, value in param_dict.items():
        placeholder = "{{" + key + "}}"
        for elem in root.xpath('//w:t|//v:t', namespaces=namespaces):
            if elem.text and placeholder.lower() in elem.text.lower():
                pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
                elem.text = re.sub(pattern, value, elem.text, flags=re.IGNORECASE)

def _process_paragraphs(paragraphs, param_dict):
    def replace_placeholders(text):
        for key, value in param_dict.items():
            pattern = r"\{\{\s*" + re.escape(key) + r"\s*\}\}"
            text = re.sub(pattern, value, text, flags=re.IGNORECASE)
        return text
    for para in paragraphs:
        if "{{" in para.text:
            full_text = "".join(run.text for run in para.runs)
            new_text = replace_placeholders(full_text)
            if new_text != full_text:
                for r in para.runs:
                    r.text = ""
                para.add_run(new_text)

def _process_cell(cell, param_dict):
    _process_paragraphs(cell.paragraphs, param_dict)
    for nested_table in getattr(cell, 'tables', []):
        for row in nested_table.rows:
            for c in row.cells:
                _process_cell(c, param_dict)

def fill_template(df: pd.DataFrame, template_path: str):
    param_dict = {str(p).strip().lower(): str(v) for p, v in zip(df["Parameters"], df["Value"])}
    doc = Document(template_path)

    _replace_in_xml(doc, param_dict)
    _process_paragraphs(doc.paragraphs, param_dict)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_cell(cell, param_dict)

    for section in doc.sections:
        _replace_in_xml(section.header, param_dict)
        _replace_in_xml(section.footer, param_dict)
        _process_paragraphs(section.header.paragraphs, param_dict)
        _process_paragraphs(section.footer.paragraphs, param_dict)
        for table in list(getattr(section.header, 'tables', [])) + list(getattr(section.footer, 'tables', [])):
            for row in table.rows:
                for cell in row.cells:
                    _process_cell(cell, param_dict)

    return doc
